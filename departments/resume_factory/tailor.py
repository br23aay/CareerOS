"""
departments/resume_factory/tailor.py — fresh, ATS-safe resume per job.

Honest by construction: it only ever REORDERS, SELECTS and RE-EMPHASISES the
real facts in core/master_cv.py against a job description. It cannot invent a
skill — if it isn't in your master CV, it can't appear.

Pipeline per job:
  1. parse JD -> keywords present in your real skill set
  2. choose a role-family lead clause for the summary
  3. order skill groups by JD relevance, surface matched keywords first
  4. pick the 2-3 most relevant projects (by tag overlap)
  5. emit an ATS-safe .docx (single column, no tables, no graphics)
  6. report an ATS keyword-coverage score

Output: storage/resumes/CV_<Company>_<role>.docx
"""

import re
import subprocess
import sys
from pathlib import Path

sys.path.append(str(Path(__file__).resolve().parent.parent.parent))
from core.base_agent import BaseAgent
from core import config, master_cv as M


def _jd_keywords(jd: str) -> set:
    """Lowercase token + bigram set from the job description."""
    words = re.findall(r"[a-zA-Z][a-zA-Z+.#-]{1,}", jd.lower())
    bigrams = {f"{words[i]} {words[i+1]}" for i in range(len(words) - 1)}
    return set(words) | bigrams


def _all_skill_terms() -> list:
    terms = []
    for group in M.SKILL_GROUPS.values():
        terms += [s.lower() for s in group]
    return terms


def choose_lead(jd: str) -> str:
    t = jd.lower()
    if any(k in t for k in ("robot", "mujoco", "manipulation", "control")):
        return M.LEAD_CLAUSES["robotics"]
    if any(k in t for k in ("llm", "rag", "prompt", "language model", "genai")):
        return M.LEAD_CLAUSES["llm"]
    if any(k in t for k in ("research", "publication", "novel")):
        return M.LEAD_CLAUSES["research"]
    if any(k in t for k in ("vision", "image", "detection", "ocr")):
        return M.LEAD_CLAUSES["vision"]
    if any(k in t for k in ("data scientist", "data analy", "pipeline")):
        return M.LEAD_CLAUSES["data"]
    if any(k in t for k in ("software engineer", "developer", "backend")):
        return M.LEAD_CLAUSES["software"]
    return M.LEAD_CLAUSES["llm"]


def order_skill_groups(jd_kw: set) -> list:
    """Return skill groups ordered by how many terms the JD mentions."""
    scored = []
    for name, skills in M.SKILL_GROUPS.items():
        hits = sum(1 for s in skills if s.lower() in jd_kw
                   or any(s.lower() in k for k in jd_kw))
        scored.append((hits, name, skills))
    scored.sort(key=lambda x: -x[0])
    # keep groups with at least one hit first, then the rest for completeness
    return [(n, s) for _, n, s in scored]


def pick_projects(jd_kw: set, n: int = 3) -> list:
    scored = []
    for proj in M.PROJECTS:
        hits = sum(1 for tag in proj["tags"]
                   if tag in jd_kw or any(tag in k for k in jd_kw))
        scored.append((hits, proj))
    scored.sort(key=lambda x: -x[0])
    chosen = [p for h, p in scored if h > 0][:n]
    return chosen or [M.PROJECTS[0]]  # always show at least the flagship


def ats_score(jd_kw: set) -> dict:
    """How many of your real skills the JD asks for and you cover."""
    skill_terms = _all_skill_terms()
    jd_skill_hits = [s for s in skill_terms
                     if s in jd_kw or any(s in k for k in jd_kw)]
    return {"covered": sorted(set(jd_skill_hits)),
            "count": len(set(jd_skill_hits))}


class ResumeFactory(BaseAgent):
    name = "resume_factory.tailor"

    def build(self, job: dict) -> dict:
        jd = f"{job.get('title','')} {job.get('description','')}"
        jd_kw = _jd_keywords(jd)
        lead = choose_lead(jd)
        groups = order_skill_groups(jd_kw)
        projects = pick_projects(jd_kw)
        ats = ats_score(jd_kw)
        summary = f"{lead}. {M.SUMMARY_BASE}"
        company = re.sub(r"[^A-Za-z0-9]+", "_",
                         job.get("company", "company")).strip("_")[:40]
        role = re.sub(r"[^A-Za-z0-9]+", "_",
                      job.get("title", "role")).strip("_")[:30]
        stem = f"CV_{company}_{role}"

        # Build a real ATS-safe .docx natively (no pandoc dependency).
        docx_path = self._build_docx(summary, groups, projects, stem, jd_kw)
        self.log.info(f"Tailored resume -> {Path(docx_path).name} "
                      f"(ATS keywords covered: {ats['count']})")
        return {"docx": docx_path, "ats_covered": ats["count"],
                "ats_keywords": ats["covered"], "lead": lead,
                "projects": [p["title"] for p in projects]}

    def _order_points(self, role, jd_kw):
        """Order an experience role's bullets so JD-relevant ones lead."""
        def relevance(point):
            pl = point.lower()
            return sum(1 for kw in jd_kw if kw in pl)
        # keep all points (honest — nothing removed), just reorder by relevance
        return sorted(role["points"], key=relevance, reverse=True)

    def _build_docx(self, summary, groups, projects, stem, jd_kw=None) -> str:
        """Write an ATS-safe .docx directly with python-docx (single column,
        no tables, no graphics, standard headings — clean for ATS parsers)."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        c = M.CONTACT
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

        def heading(text):
            h = doc.add_paragraph()
            run = h.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            h.space_after = Pt(2)

        def bullet(text):
            doc.add_paragraph(text, style="List Bullet")

        # header
        name = doc.add_paragraph()
        r = name.add_run(c["name"])
        r.bold = True; r.font.size = Pt(18)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact = doc.add_paragraph(
            f"{c['location']} | {c['phone']} | {c['email']}")
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links = doc.add_paragraph(
            f"{c['linkedin']} | {c['github']} | {c['portfolio']}")
        links.alignment = WD_ALIGN_PARAGRAPH.CENTER

        heading("Professional Summary")
        doc.add_paragraph(summary)

        heading("Core Technical Skills")
        for gname, skills in groups:
            if skills:
                p = doc.add_paragraph()
                p.add_run(f"{gname}: ").bold = True
                p.add_run(", ".join(skills))

        heading("Education")
        for e in M.EDUCATION:
            p = doc.add_paragraph()
            p.add_run(f"{e['title']} — {e['place']} ({e['dates']})").bold = True
            for pt in e["points"]:
                bullet(pt)

        heading("Experience")
        for x in M.EXPERIENCE:
            p = doc.add_paragraph()
            p.add_run(f"{x['title']} — {x['place']} ({x['dates']})").bold = True
            # order this role's bullets so the JD-relevant ones come first
            pts = self._order_points(x, jd_kw) if jd_kw else x["points"]
            for pt in pts:
                bullet(pt)

        heading("Key Projects (selected for this role)")
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(f"{proj['title']} ({proj['dates']})").bold = True
            for pt in proj["points"]:
                bullet(pt)

        heading("Certifications — 49 Microsoft Learn Badges")
        for ct in M.CERTIFICATIONS:
            bullet(ct)

        heading("Publication")
        doc.add_paragraph(M.PUBLICATION)

        heading("Additional Information")
        for a in M.ADDITIONAL:
            bullet(a)

        out = config.RESUMES / f"{stem}.docx"
        doc.save(str(out))
        return str(out)
